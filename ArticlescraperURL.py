#IMPORTS
from newspaper import Article
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import PIL
import requests
import re

#DECLARATIONS
imglist = list()
vdolist = list()
document = Document()
vidoecount = 0

#FUNCTIONS


#ADD A LINE OF TEXT IN THE DOCUMENT WITH FONT TYPE AND ALIGNMENT
def addline(lne,typ,algn): 
    docline = document.add_paragraph()
    f = docline.add_run(lne)
    doclinef = f.font
    if typ == "I":
        doclinef.italic = True
    elif typ == "B":
        doclinef.bold = True
    elif typ == "BI" or typ == "IB":
        doclinef.italic = True
        doclinef.bold = True
    else:
        doclinef.italic = False
        doclinef.bold = False
    if algn == "L":
        docline.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif algn == "C":
        docline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        docline.alignment = WD_ALIGN_PARAGRAPH.LEFT 
    

#DOCUMENTING TITLE
def  writetitle(ttle):
    doctitle = document.add_heading(ttle, level=1)
    doctitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

#DOCUMENTING ARTICLE
def writearticle(txt):
    doctxt = document.add_paragraph(txt)
    

#ADD IMAGES IN A TABLE
def addimages (imglist,m,r,c):
    imgstbl = document.add_table(rows = r, cols = c)
    urllist = list(imglist)
    r = 0
    c = 0
    if m == 2:
        for url in urllist :
            hdr_cells = imgstbl.cell
            paratbl = hdr_cells(r,c).paragraphs[0]
            pararun = paratbl.add_run()
            try:
                respon = requests.get(url)
                img = BytesIO(respon.content)
                pararun.add_picture(img, width=Inches(2), height = Inches(2))
                pararun.add_break()
                paratbl.add_run(url)
                c = c + 1
                if  c == 2:
                    r = r + 1
                    c = 0
            except:
                paratbl.add_run("Broken link for image")
                c = c + 1
                if  c == 2:
                    r = r + 1
                    c = 0
                
    elif m == 3:
        for url in urllist :
            hdr_cells = imgstbl.cell
            paratbl = hdr_cells(r,c).paragraphs[0]
            pararun = paratbl.add_run()
            try:
                respon = requests.get(url)
                img = BytesIO(respon.content)
                pararun.add_picture(img, width=Inches(2), height = Inches(2))
                pararun.add_break()
                paratbl.add_run(imgurl)
                c = c + 1
                if  c == 3:
                    r = r + 1
                    c = 0
            except:
                paratbl.add_run(url)
                pararun.add_break()
                paratbl.add_run("Broken link for image")
                c = c + 1
                if  c == 3:
                    r = r + 1
                    c = 0
    

#GETTING THE TOP IMAGE
def topimg(imgurl):
    response = requests.get(imgurl)
    img = BytesIO(response.content)  
    document.add_picture(img, width=Inches(3), height=Inches(3))
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

#DOCUMENTING KEYWORDS
def writekwrds(txt):
    addline("Keywords","B","L")
    kwrd = document.add_paragraph()
    for wrd in txt:
        w = kwrd.add_run(wrd)
        kwrdfnt = w.font
        kwrdfnt.bold = True
        kwrdfnt.italic = True
        kwrd.add_run(", ")
        
while True:
	#GETTING THE LINK
    link = input("Enter the URL to process: ")
	
	#SCRAPING THE DATA
    url = link
    try:
        article = Article(url)
        article.download()
        article.parse()
        #article.nlp()
    except:
        print(url, "Cannot Access URL")
        break
     
	#GETTING NAME OF THE DOCUMENT AND INITIATING THE PROCESS
    docname = input("\nInput a name for the document [ADD '.docx' in the end]: ")

    #WRITING THE DOCUMENT
    #TITLE
    title = article.title
    writetitle(title)
    #URL
    addline(url,"I","L")
    #TOP IMAGE
    addline("Top Image","B","L")
    topimgurl = article.top_image
    try:
        topimg(topimgurl)
    except:
        addline("Broken link for image","B","C")
    #ARTICLE
    addline("Article","B","L")
    writearticle(article.text)
    #IMAGES
    addline("Images","B","L")
    imglist = article.images
    imglstlen = len(imglist)
    if (imglstlen % 2) == 0 :
        noofcols = 2
        noofrows = int(imglstlen/2)
        addimages(imglist,2,noofrows,noofcols)
    elif (imglstlen % 3) == 0 :
        noofcols = 3
        noofrows = int(imglstlen/3)
        addimages(imglist,3,noofrows,noofcols)
    else:
        div = int(imglstlen / 3) + 1
        noofcols = 3
        addimages(imglist,3,div,noofcols)
    #VIDEO LINKS
    vdolist = article.movies
    addline("Videos","B","L")
    if len(vdolist) > 0 :
        for vdo in vdolist:
            vidoecount = vidoecount + 1
            addline("Article Videos "+str(vidoecount)+":","B","L")
            addline(vdo,"I","L")
            vidoecount = 0
    else:
        addline("No Videos for this article.","B","C")
    #SAVING DOCUMENT
    document.add_page_break()
    document.save(docname)
    print (url + "\nThe above URL has been processed and the article is saved in " + docname)