#IMPORTS
from newspaper import Article
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import PIL
import requests
import re

#DECLARATIONS
urlregex = "http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\(\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+"
linkslist = #IMPORTS
from newspaper import Article
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import PIL
import requests
import re

#DECLARATIONS
urlregex = "http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\(\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+"
linkslist = list()
selectedlinks = list()
titlelist = list()
imglist = list()
vdolist = list()
document = Document()
Articlecount = 0
vidoecount = 0

#FUNCTIONS

#PRINTLINKS FUNCTION
def printlinks(strt,stp):
    print("\nThe URLs to process are as follows:")
    Linkcount = 0
    for l in linkslist[strt:stp]:
        Linkcount = Linkcount + 1
        print(Linkcount, l)
        

#ADD A LINE OF TEXT IN THE DOCUMENT WITH FONT TYPE AND ALIGNMENT
def addline(lne,typ,algn): 
    docline = document.add_paragraph()
    f = docline.add_run(lne)
    doclinef = f.font
    
    if typ == "I":
        doclinef.italic = True
    elif typ == "B":
        doclinef.bold = True
    elif typ == "BI" or typ == "IB":
        doclinef.italic = True
        doclinef.bold = True
    else:
        doclinef.italic = False
        doclinef.bold = False
    if algn == "L":
        docline.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif algn == "C":
        docline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        docline.alignment = WD_ALIGN_PARAGRAPH.LEFT 
    

#DOCUMENTING TITLE
def  writetitle(ttle):
    doctitle = document.add_heading(ttle, level=1)
    doctitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

#DOCUMENTING ARTICLE
def writearticle(txt):
    doctxt = document.add_paragraph(txt)

#ADD IMAGES IN A TABLE
def addimages (imglist,m,r,c):
    imgstbl = document.add_table(rows = r, cols = c)
    urllist = list(imglist)
    r = 0
    c = 0
    if m == 2:
        for url in urllist :
            hdr_cells = imgstbl.cell
            paratbl = hdr_cells(r,c).paragraphs[0]
            pararun = paratbl.add_run()
            try:
                respon = requests.get(url)
                img = BytesIO(respon.content)
                pararun.add_picture(img, width=Inches(2), height = Inches(2))
                pararun.add_break()
                paratbl.add_run(url)
                c = c + 1
                if  c == 2:
                    r = r + 1
                    c = 0
            except:
                paratbl.add_run("Broken link for image")
                c = c + 1
                if  c == 2:
                    r = r + 1
                    c = 0
                
    elif m == 3:
        for url in urllist :
            hdr_cells = imgstbl.cell
            paratbl = hdr_cells(r,c).paragraphs[0]
            pararun = paratbl.add_run()
            try:
                respon = requests.get(url)
                img = BytesIO(respon.content)
                pararun.add_picture(img, width=Inches(2), height = Inches(2))
                pararun.add_break()
                paratbl.add_run(imgurl)
                c = c + 1
                if  c == 3:
                    r = r + 1
                    c = 0
            except:
                paratbl.add_run(url)
                pararun.add_break()
                paratbl.add_run("Broken link for image")
                c = c + 1
                if  c == 3:
                    r = r + 1
                    c = 0
    
#GETTING THE TOP IMAGE
def topimg(imgurl):
    response = requests.get(imgurl)
    img = BytesIO(response.content)  
    document.add_picture(img, width=Inches(3), height=Inches(3))
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

#DOCUMENTING KEYWORDS
def writekwrds(txt):
    addline("Keywords","B","L")
    kwrd = document.add_paragraph()
    for wrd in txt:
        w = kwrd.add_run(wrd)
        kwrdfnt = w.font
        kwrdfnt.bold = True
        kwrdfnt.italic = True
        kwrd.add_run(", ")

#INTRODUCTION
Introtxt="""
            ______________________________________________________________
            |                   THE ARTICLE SCRAPER V2                 |
            |                      AUTHOR: SATHISH K V                   |
            |------------------------------------------------------------|
            | This program allows the user to scrape the content from an |
            | article in the internet using its URL alone and convert the|
            | content into a MS-Word Document. This program uses the     |
            | 'newspaper3k' and 'Python-docx' library. Be sure to load   |
            | them before using this program. This prorgam gives the     |
            | following from an article in the Word Document,            |
            | 1. The Title and the URL                                   |
            | 2. Content                                                 |
            | 3. Top Image of the Article                                |
            | 4. Summary of the Article                                  |
            | 5. Keywords from the Article                               |
            |------------------------------------------------------------|
            | INSTRUCTIONS:                                              |
            | 1. Collect the URLs of the articles you want to obtain and |
            |    list them one URL per line in a Text file (*.txt).      |
            | 2. Now, initiate the program and follow the instructions,  |
            |    Provide the location or the name of the text file that  |
            |    contains the URLs.                                      |
            | 3. Select a set of links to process or select all depending|
            |    upon your need. Larger number of URLs (i.e. >30) might  |
            |    take longer. So, it is ideal to split and process.      |
            | 4. Provide a valid name for your Word Document with the    |
            |    '*.docx' extension.                                     | 
            | 5. Once the process is completed, the title of all the     |
            |    articles processed will be listed, marking the end of   |
            |    the process.                                            |
            |____________________________________________________________|
            |                                                            |
            | NOTE: If you are running the program for the first time,   |
            | Run the 'prereqs.py' file before executing the program     |
            |____________________________________________________________|"""
            


print(Introtxt)

while True:

    #GETTING THE FILE NAME WITH LINKS
    linksfile = input("Enter the name of the text file that contains the links or enter the full file location: \n")
    try:
        fileh = open(linksfile)
        for line in fileh :
            line = line.strip()
            links = re.findall(urlregex,line)
            for link in links :
                linkslist.append(link)
        totlinksinfile = len(linkslist)
        print("Number of links found in the file:", totlinksinfile)
    except:
        print("An Error occured, please check the file name/location")
        continue

    #GETTING NUMBER OF LINKS TO PROCESS
    noflinks2process = input("\nHow many links from the first would you like to process? \n[Use N or N-N Format or Use 'A' for All Links(Time Consuming, More Computation)] \n")
    n2nfrmt = noflinks2process.find("-")
    if n2nfrmt == -1 :
        if noflinks2process == 'A':
            nofl2pstrt = 0
            nofl2pstp = len(linkslist)
            print("The number of links to be processed is" ,nofl2pstp)
            printlinks(0,nofl2pstp)
        elif int(noflinks2process) < totlinksinfile:
            nofl2pstrt = 0
            nofl2pstp = int(noflinks2process)
            print("The number of links to be processed is" ,nofl2pstp)
            printlinks(0,nofl2pstp)
        elif int(noflinks2process) == len(linkslist):
            nofl2pstrt = 0
            nofl2pstp = int(noflinks2process)
            print("The number of links to be processed is" ,nofl2pstp)
            printlinks(0,nofl2pstp)
        else:    
            print("ERROR: The number of links provided are greater than the number of links in the file.")
            continue
    else:
        nofl2pstrt = int(noflinks2process[0:n2nfrmt])
        nofl2pstp = int(noflinks2process[n2nfrmt+1:len(noflinks2process)])
        if nofl2pstp < nofl2pstrt:
            print("ERROR: The number of links to stop the process is less than the number of link to start the process.")
            continue
        elif nofl2pstrt == 1:
            nofl2pstrt = 0
            print("Start of the link to process is", int(nofl2pstrt))
            print("End of the link to process is", int(nofl2pstp))
            printlinks(nofl2pstrt,nofl2pstp)
        else:    
            print("Start of the link to process is", int(nofl2pstrt))
            print("End of the link to process is", int(nofl2pstp))
            printlinks(nofl2pstrt,nofl2pstp)

    #GETTING NAME OF THE DOCUMENT AND INITIATING THE PROCESS
    docname = input("\nInput a name for the document [ADD '.docx' in the end]: ")
    selectedlinks = linkslist[nofl2pstrt:nofl2pstp]
    print("\nThe article scraping operation has been completed for the following titles:")
    for link in selectedlinks:
        #SCRAPING THE DATA
        url = link
        try:
            article = Article(url)
            article.download()
            article.parse()
            article.nlp()
        except:
            print(url, "Cannot Access URL")
        titlelist.append(article.title)
        #WRITING THE DOCUMENT
        #TITLE
        title = article.title
        writetitle(title)
        #URL
        addline(url,"I","L")
        #ARTICLE
        writearticle(article.text)
        #IMAGES
        imglist = article.images
        imglstlen = len(imglist)
        if (imglstlen % 2) == 0 :
            noofcols = 2
            noofrows = int(imglstlen/2)
            addimages(imglist,2,noofrows,noofcols)
        elif (imglstlen % 3) == 0 :
            noofcols = 3
            noofrows = int(imglstlen/3)
            addimages(imglist,3,noofrows,noofcols)
        else:
            div = int(imglstlen / 3) + 1
            noofcols = 3
            addimages(imglist,3,div,noofcols)
        #VIDEO LINKS
        vdolist = article.movies
        addline("Videos","B","L")
        if len(vdolist) > 0 :
            for vdo in vdolist:
                vidoecount = vidoecount + 1
                addline("Article Videos "+str(vidoecount)+":","B","L")
                addline(vdo,"I","L")
            vidoecount = 0
        else:
            addline("No Videos for this article.","B","C")
        #SUMMARY
        summry = article.summary
        addline("Summary","B","L")
        writearticle(summry)
        #KEYWORDS
        kwrds = article.keywords
        writekwrds(kwrds)
        #SAVING DOCUMENT
        document.add_page_break()
        document.save(docname)
        Articlecount = Articlecount + 1
        print(Articlecount, title)

    list()
selectedlinks = list()
titlelist = list()
imglist = list()
vdolist = list()
document = Document()
Articlecount = 0
vidoecount = 0

#FUNCTIONS

#PRINTLINKS FUNCTION
def printlinks(strt,stp):
    print("\nThe URLs to process are as follows:")
    Linkcount = 0
    for l in linkslist[strt:stp]:
        Linkcount = Linkcount + 1
        print(Linkcount, l)
        

#ADD A LINE OF TEXT IN THE DOCUMENT WITH FONT TYPE AND ALIGNMENT
def addline(lne,typ,algn): 
    docline = document.add_paragraph()
    f = docline.add_run(lne)
    doclinef = f.font
    
    if typ == "I":
        doclinef.italic = True
    elif typ == "B":
        doclinef.bold = True
    elif typ == "BI" or typ == "IB":
        doclinef.italic = True
        doclinef.bold = True
    else:
        doclinef.italic = False
        doclinef.bold = False
    if algn == "L":
        docline.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif algn == "C":
        docline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        docline.alignment = WD_ALIGN_PARAGRAPH.LEFT 
    

#DOCUMENTING TITLE
def  writetitle(ttle):
    doctitle = document.add_heading(ttle, level=1)
    doctitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

#DOCUMENTING ARTICLE
def writearticle(txt):
    doctxt = document.add_paragraph(txt)
    
#GETTING THE TOP IMAGE
def writeimg(imgurl):
    response = requests.get(imgurl)
    img = BytesIO(response.content)  
    document.add_picture(img, width=Inches(3), height=Inches(3))
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

#DOCUMENTING KEYWORDS
def writekwrds(txt):
    addline("Keywords","B","L")
    kwrd = document.add_paragraph()
    for wrd in txt:
        w = kwrd.add_run(wrd)
        kwrdfnt = w.font
        kwrdfnt.bold = True
        kwrdfnt.italic = True
        kwrd.add_run(", ")

#INTRODUCTION
Introtxt="""
            ______________________________________________________________
            |                   THE ARTICLE SCRAPER V1.1                 |
            |                      AUTHOR: SATHISH K V                   |
            |------------------------------------------------------------|
            | This program allows the user to scrape the content from an |
            | article in the internet using its URL alone and convert the|
            | content into a MS-Word Document. This program uses the     |
            | 'newspaper3k' and 'Python-docx' library. Be sure to load   |
            | them before using this program. This prorgam gives the     |
            | following from an article in the Word Document,            |
            | 1. The Title and the URL                                   |
            | 2. Content                                                 |
            | 3. Top Image of the Article                                |
            | 4. Summary of the Article                                  |
            | 5. Keywords from the Article                               |
            |------------------------------------------------------------|
            | INSTRUCTIONS:                                              |
            | 1. Collect the URLs of the articles you want to obtain and |
            |    list them one URL per line in a Text file (*.txt).      |
            | 2. Now, initiate the program and follow the instructions,  |
            |    Provide the location or the name of the text file that  |
            |    contains the URLs.                                      |
            | 3. Select a set of links to process or select all depending|
            |    upon your need. Larger number of URLs (i.e. >30) might  |
            |    take longer. So, it is ideal to split and process.      |
            | 4. Provide a valid name for your Word Document with the    |
            |    '*.docx' extension.                                     | 
            | 5. Once the process is completed, the title of all the     |
            |    articles processed will be listed, marking the end of   |
            |    the process.                                            |
            |____________________________________________________________|
            |                                                            |
            | NOTE: If you are running the program for the first time,   |
            | Run the 'prereqs.py' file before executing the program     |
            |____________________________________________________________|"""
            


print(Introtxt)

#GETTING THE FILE NAME WITH LINKS
linksfile = input("Enter the name of the text file that contains the links or enter the full file location: \n")
try:
    fileh = open(linksfile)
    for line in fileh :
        line = line.strip()
        links = re.findall(urlregex,line)
        for link in links :
            linkslist.append(link)
    totlinksinfile = len(linkslist)
    print("Number of links found in the file:", totlinksinfile)
except:
    print("An Error occured, please check the file name/location")

#GETTING NUMBER OF LINKS TO PROCESS
noflinks2process = input("\nHow many links from the first would you like to process? \n[Use N or N-N Format or Use 'A' for All Links(Time Consuming, More Computation)] \n")
n2nfrmt = noflinks2process.find("-")
if n2nfrmt == -1 :
    if noflinks2process == 'A':
        nofl2pstrt = 0
        nofl2pstp = len(linkslist)
        print("The number of links to be processed is" ,nofl2pstp)
        printlinks(0,nofl2pstp)
    elif int(noflinks2process) < totlinksinfile:
        nofl2pstrt = 0
        nofl2pstp = int(noflinks2process)
        print("The number of links to be processed is" ,nofl2pstp)
        printlinks(0,nofl2pstp)
    elif int(noflinks2process) == len(linkslist):
        nofl2pstrt = 0
        nofl2pstp = int(noflinks2process)
        print("The number of links to be processed is" ,nofl2pstp)
        printlinks(0,nofl2pstp)
    else:    
        print("ERROR: The number of links provided are greater than the number of links in the file.")   
else:
    nofl2pstrt = int(noflinks2process[0:n2nfrmt])
    nofl2pstp = int(noflinks2process[n2nfrmt+1:len(noflinks2process)])
    if nofl2pstp < nofl2pstrt:
         print("ERROR: The number of links to stop the process is less than the number of link to start the process.")
    elif nofl2pstrt == 1:
        nofl2pstrt = 0
        print("Start of the link to process is", int(nofl2pstrt))
        print("End of the link to process is", int(nofl2pstp))
        printlinks(nofl2pstrt,nofl2pstp)
    else:    
        print("Start of the link to process is", int(nofl2pstrt))
        print("End of the link to process is", int(nofl2pstp))
        printlinks(nofl2pstrt,nofl2pstp)

#GETTING NAME OF THE DOCUMENT AND INITIATING THE PROCESS
docname = input("\nInput a name for the document [ADD '.docx' in the end]: ")
selectedlinks = linkslist[nofl2pstrt:nofl2pstp]
print("\nThe article scraping operation has been completed for the following titles:")
for link in selectedlinks:
#SCRAPING THE DATA
    url = link
    try:
        article = Article(url)
        article.download()
        article.parse()
        article.nlp()
    except:
        print(url, "Cannot Access URL")
        continue
    titlelist.append(article.title)
#WRITING THE DOCUMENT
    #TITLE
    title = article.title
    writetitle(title)
    #URL
    addline(url,"I","L")
    #ARTICLE
    writearticle(article.text)
    #IMAGES
    imglist = article.images
    if len(imglist)>0 :
        for im in imglist:
            try:
                writeimg(im)
                #writearticle(arg4lnk)
            except:
                addline("Broken Link for the Image.","B","C")
    else:
        addline("No Image for this article.","B","C")
    #VIDEO LINKS
    vdolist = article.movies
    addline("Videos","B","L")
    if len(vdolist) > 0 :
        for vdo in vdolist:
            vidoecount = vidoecount + 1
            addline("Article Videos "+str(vidoecount)+":","B","L")
            addline(vdo,"I","L")
        vidoecount = 0
    else:
        addline("No Videos for this article.","B","C")
    #SUMMARY
    summry = article.summary
    addline("Summary","B","L")
    writearticle(summry)
    #KEYWORDS
    kwrds = article.keywords
    writekwrds(kwrds)
    #SAVING DOCUMENT
    document.add_page_break()
    document.save(docname)
    Articlecount = Articlecount + 1
    print(Articlecount, title)
